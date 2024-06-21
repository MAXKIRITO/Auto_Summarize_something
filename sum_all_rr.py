import os
import sys
import queue
import sounddevice as sd
import torch
import torchaudio
import json
import threading
import logging
import cv2
import numpy as np
import time
from PIL import Image
from datetime import datetime, timedelta
from docx import Document
from docx.shared import Inches
from collections import deque
import torchvision.transforms as transforms
import pyaudio
from transformers import T5Tokenizer, T5ForConditionalGeneration

# Initialize logging and directories
if not os.path.exists('log'):
    os.makedirs('log')

log_filename = datetime.utcnow().strftime('log/transcription_%Y-%m-%dT%H-%M-%SSZ.log')
logging.basicConfig(filename=log_filename, level=logging.INFO, format='%(asctime)s - %(message)s')

# Initialize queues and other shared resources
q = queue.Queue()
image_queue = queue.Queue()
transcriptions = []
summaries = []
recent_features = deque(maxlen=10)
similarity_threshold = 90
summary_interval = 60  # summary interval in seconds
last_summary_time = datetime.utcnow()

# Initialize the image model
image_model = torch.hub.load('pytorch/vision:v0.10.0', 'resnet50', pretrained=True)
image_model.eval()
if torch.cuda.is_available():
    image_model = image_model.to('cuda')

preprocess = transforms.Compose([
    transforms.Resize((224, 224)),
    transforms.ToTensor(),
    transforms.Normalize(mean=[0.485, 0.456, 0.406], std=[0.229, 0.224, 0.225]),
])

# Initialize the speech recognition model
bundle = torchaudio.pipelines.WAV2VEC2_ASR_BASE_960H
model = bundle.get_model()
labels = bundle.get_labels()
if torch.cuda.is_available():
    model = model.to('cuda')

# Initialize the T5 model and tokenizer
tokenizer = T5Tokenizer.from_pretrained('t5-small')
t5_model = T5ForConditionalGeneration.from_pretrained('t5-small')
device = torch.device('cuda' if torch.cuda.is_available() else 'cpu')
t5_model.to(device)

def list_microphones():
    p = pyaudio.PyAudio()
    info = p.get_host_api_info_by_index(0)
    num_devices = info.get('deviceCount')
    microphones = []

    for i in range(0, num_devices):
        device_info = p.get_device_info_by_host_api_device_index(0, i)
        if device_info.get('maxInputChannels') > 0:
            microphones.append((i, device_info.get('name')))

    p.terminate()
    return microphones

def select_microphone():
    microphones = list_microphones()
    for idx, (device_index, device_name) in enumerate(microphones):
        print(f"{idx}: {device_name} (Device Index: {device_index})")
    choice = int(input("Select a microphone by number: "))
    return microphones[choice][0]

def callback(indata, frames, time, status):
    if status:
        print(status, file=sys.stderr)
    q.put(bytes(indata))

def is_blurry(image, threshold=100):
    gray = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY)
    laplacian_var = cv2.Laplacian(gray, cv2.CV_64F).var()
    return laplacian_var < threshold

def extract_features(image):
    image = Image.fromarray(cv2.cvtColor(image, cv2.COLOR_BGR2RGB))
    input_tensor = preprocess(image).unsqueeze(0)
    if torch.cuda.is_available():
        input_tensor = input_tensor.to('cuda')
    with torch.no_grad():
        features = image_model(input_tensor)
    return features

def is_different(features1, features2, similarity_threshold=75):
    similarity = torch.nn.functional.cosine_similarity(features1, features2).item()
    return similarity < (similarity_threshold / 100)

def process_images(image_queue, recent_features, save_dir, similarity_threshold):
    while True:
        frame, fgmask, timestamp = image_queue.get()
        if frame is None:
            break

        blurry = is_blurry(frame)
        if not blurry:
            fg_frame = cv2.bitwise_and(frame, frame, mask=fgmask)
            frame_features = extract_features(fg_frame)
            different_from_all = all(is_different(frame_features, prev_features, similarity_threshold=similarity_threshold) for prev_features in recent_features)

            if different_from_all:
                timestamp_str = datetime.utcnow().strftime('%Y-%m-%dT%H-%M-%S.%fZ')
                save_path = os.path.join(save_dir, f'image_{timestamp_str}.jpg')
                cv2.imwrite(save_path, frame)
                recent_features.append(frame_features)
                update_word_document(transcriptions, summaries, save_dir)
                image_queue.task_done()

def greedy_decoder(emissions, labels):
    """Greedy decoder for CTC."""
    _, indices = torch.max(emissions, dim=-1)
    indices = indices.squeeze(0)
    result = []
    for i in range(indices.size(0)):
        if i > 0 and indices[i] == indices[i - 1]:
            continue
        result.append(labels[indices[i]])
    return "".join(result).replace("|", " ").replace("-", "").strip()

def transcribe_audio(waveform):
    if torch.cuda.is_available():
        waveform = waveform.to('cuda')
    with torch.no_grad():
        emissions, _ = model(waveform)
    transcription = greedy_decoder(emissions.cpu(), labels)
    return transcription

def summarize_text_t5(text):
    # 構建輸入文本
    input_text = "summarize: " + text
    # 編碼文本
    input_ids = tokenizer.encode(input_text, return_tensors='pt', max_length=512, truncation=True).to(device)
    
    # 手動設置生成過程中的一些參數
    attention_mask = torch.ones(input_ids.shape, device=device)
    
    # 設置生成配置
    generate_config = {
        "max_length": 150,
        "num_beams": 2,
        "length_penalty": 2.0,
        "early_stopping": True,
        "attention_mask": attention_mask
    }
    
    # 生成摘要
    with torch.no_grad():
        summary_ids = t5_model.generate(input_ids, **generate_config)
    
    # 解碼摘要
    summary = tokenizer.decode(summary_ids[0], skip_special_tokens=True)
    # print(f"Summary: {summary}")
    return summary

def audio_record():
    global last_summary_time

    device_index = select_microphone()

    with sd.RawInputStream(samplerate=16000, blocksize=8000, dtype='int16', channels=1, callback=callback, device=device_index):
        print('#' * 80)
        print('Press Ctrl+C to stop the recording')
        print('#' * 80)

        buffer = b""
        while True:
            data = q.get()
            buffer += data
            if len(buffer) > 16000 * 2:  # Collect 1 second of audio
                audio_data = torch.FloatTensor(np.frombuffer(buffer, dtype=np.int16).astype(np.float32) / 32768.0).unsqueeze(0)
                if torch.cuda.is_available():
                    audio_data = audio_data.to('cuda')
                transcription = transcribe_audio(audio_data)
                print(f"Transcription: {transcription}")
                timestamp_str = datetime.utcnow().strftime('%Y-%m-%dT%H-%M-%S.%fZ')
                transcriptions.append((timestamp_str, transcription))
                update_word_document(transcriptions, summaries, 'img_a')
                buffer = b""

            # Check if it's time to summarize
            if (datetime.utcnow() - last_summary_time).total_seconds() >= summary_interval:
                recent_texts = ' '.join([t[1] for t in transcriptions[-10:]])  # use last 10 transcriptions
                summary_text = summarize_text_t5(recent_texts)
                summary_timestamp = datetime.utcnow().strftime('%Y-%m-%dT%H-%M-%S.%fZ')
                summaries.append((summary_timestamp, summary_text))
                last_summary_time = datetime.utcnow()
                update_word_document(transcriptions, summaries, 'img_a')

def capture_images(image_queue, capture_interval=1):
    cap = cv2.VideoCapture(2)  # 攝像頭索引，根據需要更改
    if not cap.isOpened():
        print("Error: Could not open camera.")
        return

    # 嘗試設置攝像頭最高解析度
    cap.set(cv2.CAP_PROP_FRAME_WIDTH, 1920)
    cap.set(cv2.CAP_PROP_FRAME_HEIGHT, 1080)
    
    # 獲取並打印當前解析度設置
    width = int(cap.get(cv2.CAP_PROP_FRAME_WIDTH))
    height = int(cap.get(cv2.CAP_PROP_FRAME_HEIGHT))
    print(f"Current camera resolution: {width}x{height}")

    fgbg = cv2.createBackgroundSubtractorMOG2()
    last_saved_time = time.time()

    while True:
        ret, frame = cap.read()
        if not ret:
            break

        fgmask = fgbg.apply(frame)
        non_zero_count = np.count_nonzero(fgmask)

        if non_zero_count > (frame.shape[0] * frame.shape[1] * 0.05):
            current_time = time.time()
            if current_time - last_saved_time >= capture_interval:
                timestamp_str = datetime.utcnow().strftime('%Y-%m-%dT%H-%M-%S.%fZ')
                image_queue.put((frame, fgmask, timestamp_str))
                last_saved_time = current_time

        cv2.imshow('Video', frame)
        cv2.imshow('Foreground Mask', fgmask)

        if cv2.waitKey(1) & 0xFF == ord('q'):
            break

    image_queue.put((None, None, None))
    cap.release()
    cv2.destroyAllWindows()

def create_word_document(transcriptions, summaries, image_folder):
    doc = Document()
    doc.add_heading('Merged Transcriptions and Images', 0)

    combined_entries = []

    for timestamp, transcription in transcriptions:
        combined_entries.append((timestamp, 'transcription', transcription))

    for img_file in sorted(os.listdir(image_folder)):
        if img_file.endswith('.jpg'):
            timestamp = img_file.split('_')[1].split('.')[0]
            combined_entries.append((timestamp, 'image', os.path.join(image_folder, img_file)))

    combined_entries.sort()

    for entry in combined_entries:
        if entry[1] == 'transcription':
            doc.add_heading(f'Transcription at {entry[0]}', level=1)
            doc.add_paragraph(entry[2])
        elif entry[1] == 'image':
            doc.add_heading(f'Image at {entry[0]}', level=1)
            doc.add_picture(entry[2], width=Inches(4.25))

    # Append summaries at the end
    doc.add_heading('Summaries', level=1)
    for timestamp, summary in summaries:
        doc.add_heading(f'Summary at {timestamp}', level=2)
        doc.add_paragraph(summary)

    doc.save('merged_transcriptions_and_images.docx')

def update_word_document(transcriptions, summaries, image_folder):
    create_word_document(transcriptions, summaries, image_folder)

def main():
    # Start the image processing thread
    save_dir = 'img_a'
    if not os.path.exists(save_dir):
        os.makedirs(save_dir)

    processing_thread = threading.Thread(target=process_images, args=(image_queue, recent_features, save_dir, similarity_threshold))
    processing_thread.start()

    # Start the audio recording thread
    audio_thread = threading.Thread(target=audio_record)
    audio_thread.start()

    # Start the image capture thread
    capture_thread = threading.Thread(target=capture_images, args=(image_queue,))
    capture_thread.start()

    # Join threads
    capture_thread.join()
    processing_thread.join()
    audio_thread.join()

if __name__ == "__main__":
    main()
